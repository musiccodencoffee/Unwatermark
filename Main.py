
import re
import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog
import os
try:
    import docx
except ImportError:
    docx = None

# Liste der bekannten Wasserzeichen-Zeichen (Unicode)
WATERMARK_CHARS = [
    '\u00A0',  # No-Break Space (U+00A0)
    '\u00AC',  # Not Sign (U+00AC)
    '\u00AD',  # Soft Hyphen (U+00AD)
    '\u2004',  # Three-Per-Em Space (U+2004)
    '\u2005',  # Four-Per-Em Space (U+2005)
    '\u2006',  # Six-Per-Em Space (U+2006)
    '\u2007',  # Figure Space (U+2007)
    '\u2008',  # Punctuation Space (U+2008)
    '\u2009',  # Thin Space (U+2009)
    '\u200A',  # Hair Space (U+200A)
    '\u202F',  # Narrow No-Break Space (U+202F)
    '\u205F',  # Medium Mathematical Space (U+205F)
    '\u2028',  # Line Separator (U+2028)
    '\u2029',  # Paragraph Separator (U+2029)
    '\u200B',  # Zero Width Space (U+200B)
    '\u200C',  # Zero Width Non-Joiner (U+200C)
    '\u200D',  # Zero Width Joiner (U+200D)
    '\u200E',  # Left-to-Right Mark (U+200E)
    '\u200F',  # Right-to-Left Mark (U+200F)
    '\u2060',  # Word Joiner (U+2060)
    '\u2061',  # Invisible Operator (U+2061)
    '\u2062',  # Invisible Times (U+2062)
    '\u2063',  # Invisible Separator (U+2063)
    '\u2064',  # Invisible Plus (U+2064)
    '\uFEFF',  # Zero Width No-Break Space (U+FEFF)
    '\u180E',  # Mongolian Vowel Separator (deprecated, U+180E)
    '\u202A',  # Left-to-Right Embedding (U+202A)
    '\u202B',  # Right-to-Left Embedding (U+202B)
    '\u202C',  # Pop Directional Formatting (U+202C)
    '\u202D',  # Left-to-Right Override (U+202D)
    '\u202E',  # Right-to-Left Override (U+202E)
    '\u201D',  # Right Double Quotation Mark (U+201D)
    # Tag Spaces (U+E0000 bis U+E007F)
]

def contains_tag_space_or_math_op(text):
    for c in text:
        code = ord(c)
        if 0xE0000 <= code <= 0xE007F or 0x2061 <= code <= 0x2064:
            return True
    return False

def find_watermarks(text):
    """Zeigt alle Wasserzeichen-Zeichen im Text an."""
    found = []
    for char in WATERMARK_CHARS:
        if char in text:
            found.append((char, f"U+{ord(char):04X}"))
    return found

def remove_watermarks(text):
    """Entfernt alle Wasserzeichen-Zeichen aus dem Text."""
    pattern = '[' + ''.join(WATERMARK_CHARS) + ']'
    return re.sub(pattern, '', text)

def show_watermarks():
    text = input_text.get("1.0", tk.END)
    input_text.tag_remove('watermark', "1.0", tk.END)
    indices = []
    for i, c in enumerate(text):
        if c in WATERMARK_CHARS:
            indices.extend([i-1, i, i+1])
    indices = sorted(set([i for i in indices if 0 <= i < len(text)-1]))
    for i in indices:
        start = f"1.0+{i}c"
        end = f"1.0+{i+1}c"
        input_text.tag_add('watermark', start, end)
    input_text.tag_configure('watermark', background='#ffcccc')
    watermarks = find_watermarks(text)
    if watermarks:
        result = "Gefundene Wasserzeichen:\n"
        for char, code in watermarks:
            result += f"'{char}' (Unicode: {code})\n"
    else:
        result = "Keine Wasserzeichen gefunden."
    output_text.config(state='normal')
    output_text.delete("1.0", tk.END)
    output_text.insert(tk.END, result)
    output_text.config(state='disabled')

def clean_text():
    text = input_text.get("1.0", tk.END)
    cleaned = remove_watermarks(text)
    output_text.config(state='normal')
    output_text.delete("1.0", tk.END)
    output_text.insert(tk.END, cleaned)
    output_text.config(state='disabled')

def paste_clipboard():
    try:
        clipboard = root.clipboard_get()
        input_text.delete("1.0", tk.END)
        input_text.insert(tk.END, clipboard)
    except tk.TclError:
        messagebox.showerror("Fehler", "Kein Text im Clipboard gefunden.")

def copy_output():
    output = output_text.get("1.0", tk.END).strip()
    if output:
        root.clipboard_clear()
        root.clipboard_append(output)
        messagebox.showinfo("Kopiert", "Ergebnis wurde ins Clipboard kopiert.")
    else:
        messagebox.showwarning("Hinweis", "Keine Ausgabe zum Kopieren.")

def clean_clipboard():
    try:
        clipboard = root.clipboard_get()
        removed_chars = set()
        for c in clipboard:
            if c in WATERMARK_CHARS:
                removed_chars.add(c)
        cleaned = remove_watermarks(clipboard)
        root.clipboard_clear()
        root.clipboard_append(cleaned)
        if removed_chars:
            chars_info = '\n'.join([f"'{c}' (U+{ord(c):04X})" for c in removed_chars])
            messagebox.showinfo("Clipboard gereinigt", f"Das Clipboard wurde von Wasserzeichen bereinigt.\nEntfernte Zeichen:\n{chars_info}")
        else:
            messagebox.showinfo("Clipboard gereinigt", "Das Clipboard wurde bereinigt. Keine Wasserzeichen gefunden.")
    except tk.TclError:
        messagebox.showerror("Fehler", "Kein Text im Clipboard gefunden.")

def clean_file_dialog():
    file_path = filedialog.askopenfilename(title="Datei auswählen", filetypes=[("Textdateien", "*.txt"), ("Word-Dokumente", "*.docx")])
    if not file_path:
        return
    ext = os.path.splitext(file_path)[1].lower()
    try:
        removed_chars = set()
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            for c in content:
                if c in WATERMARK_CHARS:
                    removed_chars.add(c)
            cleaned = remove_watermarks(content)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(cleaned)
            if removed_chars:
                chars_info = '\n'.join([f"'{c}' (U+{ord(c):04X})" for c in removed_chars])
                messagebox.showinfo("Fertig", f"Textdatei wurde gereinigt: {os.path.basename(file_path)}\nEntfernte Zeichen:\n{chars_info}")
            else:
                messagebox.showinfo("Fertig", f"Textdatei wurde gereinigt: {os.path.basename(file_path)}\nKeine Wasserzeichen gefunden.")
        elif ext == ".docx":
            if docx is None:
                messagebox.showerror("Fehler", "Das Paket 'python-docx' ist nicht installiert.")
                return
            doc = docx.Document(file_path)
            # Absätze bereinigen
            for para in doc.paragraphs:
                for c in para.text:
                    if c in WATERMARK_CHARS:
                        removed_chars.add(c)
                para.text = remove_watermarks(para.text)
            # Tabellenzellen bereinigen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for c in cell.text:
                            if c in WATERMARK_CHARS:
                                removed_chars.add(c)
                        cell.text = remove_watermarks(cell.text)
            doc.save(file_path)
            if removed_chars:
                chars_info = '\n'.join([f"'{c}' (U+{ord(c):04X})" for c in removed_chars])
                messagebox.showinfo("Fertig", f"Word-Dokument wurde gereinigt: {os.path.basename(file_path)}\nEntfernte Zeichen:\n{chars_info}")
            else:
                messagebox.showinfo("Fertig", f"Word-Dokument wurde gereinigt: {os.path.basename(file_path)}\nKeine Wasserzeichen gefunden.")
        else:
            messagebox.showerror("Fehler", "Nur .txt und .docx werden unterstützt.")
    except Exception as e:
        messagebox.showerror("Fehler", f"Fehler beim Reinigen: {e}")
if __name__ == "__main__":
    root = tk.Tk()
    root.title("AI-Watermarks Remover")

    # Farben und Fonts
    bg_main = "#f4f6fb"
    bg_entry = "#ffffff"
    bg_btn = "#4a90e2"
    fg_btn = "#ffffff"
    font_main = ("Segoe UI", 11)
    font_btn = ("Segoe UI", 10, "bold")
    font_label = ("Segoe UI", 12, "bold")

    root.configure(bg=bg_main)
    root.grid_rowconfigure(1, weight=1)
    root.grid_rowconfigure(3, weight=1)
    root.grid_columnconfigure(0, weight=1)

    tk.Label(root, text="Text eingeben:", font=font_label, bg=bg_main).grid(row=0, column=0, sticky='w', padx=14, pady=(14,0))
    input_text = scrolledtext.ScrolledText(root, width=60, height=8, font=font_main, bg=bg_entry, relief="groove", borderwidth=2)
    input_text.grid(row=1, column=0, sticky='nsew', padx=14, pady=7)

    tk.Label(root, text="Ergebnis:", font=font_label, bg=bg_main).grid(row=3, column=0, sticky='w', padx=14, pady=(14,0))
    output_text = scrolledtext.ScrolledText(root, width=60, height=8, font=font_main, bg=bg_entry, relief="groove", borderwidth=2, state='disabled')
    output_text.grid(row=4, column=0, sticky='nsew', padx=14, pady=(7,14))

    btn_frame = tk.Frame(root, bg=bg_main)
    btn_frame.grid(row=2, column=0, sticky='ew', pady=7, padx=14)
    for i in range(6):
        btn_frame.grid_columnconfigure(i, weight=1)
    style_btn = dict(bg=bg_btn, fg=fg_btn, font=font_btn, relief="flat", activebackground="#357abd", activeforeground="#ffffff", borderwidth=0, padx=8, pady=6)
    style_btn_clean = dict(bg="#e94f4f", fg=fg_btn, font=font_btn, relief="flat", activebackground="#c0392b", activeforeground="#ffffff", borderwidth=0, padx=8, pady=6)
    style_btn_file = dict(bg="#4fc97a", fg=fg_btn, font=font_btn, relief="flat", activebackground="#388e3c", activeforeground="#ffffff", borderwidth=0, padx=8, pady=6)
    btns = []
    btns.append(tk.Button(btn_frame, text="Clipboard einfügen", command=paste_clipboard, **style_btn))
    btns.append(tk.Button(btn_frame, text="Wasserzeichen anzeigen", command=show_watermarks, **style_btn))
    btns.append(tk.Button(btn_frame, text="Wasserzeichen entfernen", command=clean_text, **style_btn))
    btns.append(tk.Button(btn_frame, text="Ergebnis ins Clipboard", command=copy_output, **style_btn))
    btns.append(tk.Button(btn_frame, text="Clipboard reinigen", command=clean_clipboard, **style_btn_clean))
    btns.append(tk.Button(btn_frame, text="Datei reinigen", command=lambda: clean_file_dialog(), **style_btn_file))
    for i, btn in enumerate(btns):
        btn.grid(row=0, column=i, sticky='ew', padx=4)
        if i == 4:
            def on_enter(e, b=btn): b.config(bg="#c0392b")
            def on_leave(e, b=btn): b.config(bg="#e94f4f")
        elif i == 5:
            def on_enter(e, b=btn): b.config(bg="#388e3c")
            def on_leave(e, b=btn): b.config(bg="#4fc97a")
        else:
            def on_enter(e, b=btn): b.config(bg="#357abd")
            def on_leave(e, b=btn): b.config(bg=bg_btn)
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)

    root.minsize(600, 400)
    root.mainloop()
    bg_entry = "#ffffff"
    bg_btn = "#4a90e2"
    fg_btn = "#ffffff"
    font_main = ("Segoe UI", 11)
    font_btn = ("Segoe UI", 10, "bold")
    font_label = ("Segoe UI", 12, "bold")

    root.configure(bg=bg_main)
    root.grid_rowconfigure(1, weight=1)
    root.grid_rowconfigure(3, weight=1)
    root.grid_columnconfigure(0, weight=1)

    tk.Label(root, text="Text eingeben:", font=font_label, bg=bg_main).grid(row=0, column=0, sticky='w', padx=14, pady=(14,0))
    input_text = scrolledtext.ScrolledText(root, width=60, height=8, font=font_main, bg=bg_entry, relief="groove", borderwidth=2)
    input_text.grid(row=1, column=0, sticky='nsew', padx=14, pady=7)

    btn_frame = tk.Frame(root, bg=bg_main)
    btn_frame.grid(row=2, column=0, sticky='ew', pady=7, padx=14)
    for i in range(6):
        btn_frame.grid_columnconfigure(i, weight=1)
    style_btn = dict(bg=bg_btn, fg=fg_btn, font=font_btn, relief="flat", activebackground="#357abd", activeforeground="#ffffff", borderwidth=0, padx=8, pady=6)
    style_btn_clean = dict(bg="#e94f4f", fg=fg_btn, font=font_btn, relief="flat", activebackground="#c0392b", activeforeground="#ffffff", borderwidth=0, padx=8, pady=6)
    btns = []
    # Neue Reihenfolge:
    btns.append(tk.Button(btn_frame, text="Clipboard einfügen", command=paste_clipboard, **style_btn))
    btns.append(tk.Button(btn_frame, text="Wasserzeichen anzeigen", command=show_watermarks, **style_btn))
    btns.append(tk.Button(btn_frame, text="Wasserzeichen entfernen", command=clean_text, **style_btn))
    btns.append(tk.Button(btn_frame, text="Ergebnis ins Clipboard", command=copy_output, **style_btn))
    btns.append(tk.Button(btn_frame, text="Clipboard reinigen", command=clean_clipboard, **style_btn_clean))
    btns.append(tk.Button(btn_frame, text="Datei reinigen", command=lambda: clean_file_dialog(), **style_btn))
    for i, btn in enumerate(btns):
        btn.grid(row=0, column=i, sticky='ew', padx=4)
        if i == 4:
            def on_enter(e, b=btn): b.config(bg="#c0392b")
            def on_leave(e, b=btn): b.config(bg="#e94f4f")
        else:
            def on_enter(e, b=btn): b.config(bg="#357abd")
            def on_leave(e, b=btn): b.config(bg=bg_btn)
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)

# Datei-Reinigungsfunktion
def clean_file_dialog():
    file_path = filedialog.askopenfilename(title="Datei auswählen", filetypes=[("Textdateien", "*.txt"), ("Word-Dokumente", "*.docx")])
    if not file_path:
        return
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            cleaned = remove_watermarks(content)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(cleaned)
            messagebox.showinfo("Fertig", f"Textdatei wurde gereinigt: {os.path.basename(file_path)}")
        elif ext == ".docx":
            if docx is None:
                messagebox.showerror("Fehler", "Das Paket 'python-docx' ist nicht installiert.")
                return
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                para.text = remove_watermarks(para.text)
            doc.save(file_path)
            messagebox.showinfo("Fertig", f"Word-Dokument wurde gereinigt: {os.path.basename(file_path)}")
        else:
            messagebox.showerror("Fehler", "Nur .txt und .docx werden unterstützt.")
    except Exception as e:
        messagebox.showerror("Fehler", f"Fehler beim Reinigen: {e}")

    tk.Label(root, text="Ergebnis:", font=font_label, bg=bg_main).grid(row=3, column=0, sticky='w', padx=14, pady=(14,0))
    output_text = scrolledtext.ScrolledText(root, width=60, height=8, font=font_main, bg=bg_entry, relief="groove", borderwidth=2, state='disabled')
    output_text.grid(row=4, column=0, sticky='nsew', padx=14, pady=(7,14))

    root.minsize(600, 400)
    root.mainloop()